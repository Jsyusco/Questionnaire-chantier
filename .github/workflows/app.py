# --- IMPORTS ET PRÉPARATION ---
import streamlit as st
import pandas as pd
import uuid
import firebase_admin
from firebase_admin import credentials, firestore
from datetime import datetime
import numpy as np
import zipfile
import io
import urllib.parse
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm # Ajout de Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ALIGN_TABLE # Ajout de WD_ALIGN_TABLE
# Importations nécessaires pour les bordures
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import io
from datetime import datetime


# --- CONFIGURATION ET STYLE (inchangés) ---
st.set_page_config(page_title="Formulaire Dynamique - Firestore", layout="centered")

st.markdown("""
<style>
    .stApp { background-color: #121212; color: #e0e0e0; }
    .main-header { background-color: #1e1e1e; padding: 20px; border-radius: 10px; margin-bottom: 20px; text-align: center; border-bottom: 3px solid #E9630C; }
    .block-container { max-width: 800px; }
    .phase-block { background-color: #1e1e1e; padding: 25px; border-radius: 12px; margin-bottom: 20px; border: 1px solid #333; }
    .question-card { background-color: transparent; padding: 15px; border-radius: 8px; margin-bottom: 15px; border-left: 3px solid #E9630C; }
    h1, h2, h3 { color: #ffffff !important; }
    .description { font-size: 0.9em; color: #EB6408; margin-bottom: 10px; }
    .mandatory { color: #F4B400; font-weight: bold; margin-left: 5px; }
    .success-box { background-color: #1e4620; padding: 15px; border-radius: 8px; border-left: 5px solid #4caf50; color: #fff; margin: 10px 0; }
    .error-box { background-color: #3d1f1f; padding: 15px; border-radius: 8px; border-left: 5px solid #ff6b6b; color: #ffdad9; margin: 10px 0; }
    .stButton > button { border-radius: 8px; font-weight: bold; padding: 0.5rem 1rem; }
    div[data-testid="stButton"] > button { width: 100%; }
</style>
""", unsafe_allow_html=True)

# --- LOGIQUE DE RENOMMAGE ET D'AFFICHAGE DU PROJET (inchangée) ---

PROJECT_RENAME_MAP = {
    'Intitulé': 'Intitulé',
    'Fournisseur Bornes AC [Bornes]': 'Fournisseur Bornes AC',
    'Fournisseur Bornes DC [Bornes]': 'Fournisseur Bornes DC',
    'L [Plan de Déploiement]': 'PDC Lent',
    'R [Plan de Déploiement]': 'PDC Rapide',
    'UR [Plan de Déploiement]': 'PDC Ultra-rapide',
    'Pré L [Plan de Déploiement]': 'PDC L pré-équipés',
    'Pré R [Plan de Déploiement]': 'PDC R pré-équipés',
    'Pré UR [Plan de Déploiement]': 'PDC UR pré-équipés',
}

DISPLAY_GROUPS = [
    ['Intitulé', 'Fournisseur Bornes AC [Bornes]', 'Fournisseur Bornes DC [Bornes]'],
    ['L [Plan de Déploiement]', 'R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    ['Pré L [Plan de Déploiement]', 'Pré R [Plan de Déploiement]','Pré UR [Plan de Déploiement]' ],
]

# -----------------------------------------------------------
# --- LOGIQUE D'ATTENTE DE PHOTOS (inchangée) ---
# -----------------------------------------------------------

SECTION_PHOTO_RULES = {
    "Bornes DC": ['R [Plan de Déploiement]', 'UR [Plan de Déploiement]'],
    "Bornes AC": ['L [Plan de Déploiement]'],
}

def get_expected_photo_count(section_name, project_data):
    if section_name not in SECTION_PHOTO_RULES:
        return None, None 

    columns = SECTION_PHOTO_RULES[section_name]
    total_expected = 0
    details = []

    for col in columns:
        val = project_data.get(col, 0)
        try:
            if pd.isna(val) or val == "":
                num = 0
            else:
                num = int(float(str(val).replace(',', '.'))) 
        except Exception:
            num = 0
        
        total_expected += num
        short_name = PROJECT_RENAME_MAP.get(col, col) 
        details.append(f"{num} {short_name}")

    detail_str = " + ".join(details)
    return total_expected, detail_str

# --- INITIALISATION FIREBASE SÉCURISÉE (inchangée) ---
def initialize_firebase():
    if not firebase_admin._apps:
        try:
            cred_dict = {
                "type": st.secrets["firebase"]["type"],
                "project_id": st.secrets["firebase"]["project_id"],
                "private_key_id": st.secrets["firebase"]["private_key_id"],
                "private_key": st.secrets["firebase"]["private_key"].replace('\\n', '\n'),
                "client_email": st.secrets["firebase"]["client_email"],
                "client_id": st.secrets["firebase"]["client_id"],
                "auth_uri": st.secrets["firebase"]["auth_uri"],
                "token_uri": st.secrets["firebase"]["token_uri"],
                "auth_provider_x509_cert_url": st.secrets["firebase"]["auth_provider_x509_cert_url"],
                "client_x509_cert_url": st.secrets["firebase"]["client_x509_cert_url"],
                "universe_domain": st.secrets["firebase"]["universe_domain"],
            }
            
            project_id = cred_dict["project_id"]
            cred = credentials.Certificate(cred_dict)
            firebase_admin.initialize_app(cred, {'projectId': project_id})
            st.sidebar.success("Connexion BDD réussie 🟢")
        
        except KeyError as e:
            st.sidebar.error(f"Erreur de configuration Secrets : Clé manquante dans la section [firebase] ({e})")
            st.stop()
        except Exception as e:
            st.sidebar.error(f"Erreur de connexion Firebase : {e}")
            st.stop()
    return firestore.client()

db = initialize_firebase()

# --- FONCTIONS DE CHARGEMENT ET SAUVEGARDE FIREBASE ---

@st.cache_data(ttl=3600)
def load_form_structure_from_firestore():
    try:
        docs = db.collection('formsquestions').order_by('id').get()
        data = [doc.to_dict() for doc in docs]
        if not data: return None
        df = pd.DataFrame(data)
        df.columns = df.columns.str.strip()
        
        rename_map = {'Conditon value': 'Condition value', 'condition value': 'Condition value', 'Condition Value': 'Condition value', 'Condition': 'Condition value', 'Conditon on': 'Condition on', 'condition on': 'Condition on'}
        actual_rename = {k: v for k, v in rename_map.items() if k in df.columns}
        df = df.rename(columns=actual_rename)
        
        expected_cols = ['options', 'Description', 'Condition value', 'Condition on', 'section', 'id', 'question', 'type', 'obligatoire']
        for col in expected_cols:
            if col not in df.columns: df[col] = np.nan 
        
        df['options'] = df['options'].fillna('')
        df['Description'] = df['Description'].fillna('')
        df['Condition value'] = df['Condition value'].fillna('')
        df['Condition on'] = df['Condition on'].apply(lambda x: int(x) if pd.notna(x) and str(x).isdigit() else 0)
        
        for col in df.select_dtypes(include=['object']).columns:
            df[col] = df[col].astype(str).str.strip()
            try:
                df[col] = df[col].apply(lambda x: x.encode('utf-8', 'ignore').decode('utf-8', 'ignore'))
            except Exception: pass 
        return df
    except Exception as e:
        return None

@st.cache_data(ttl=3600)
def load_site_data_from_firestore():
    try:
        docs = db.collection('Sites').get()
        data = [doc.to_dict() for doc in docs]
        if not data: return None
        df_site = pd.DataFrame(data)
        df_site.columns = df_site.columns.str.strip()
        return df_site
    except Exception as e:
        return None

def save_form_data(collected_data, project_data):
    """
    Sauvegarde les données dans Firestore.
    Pour les fichiers, on ne sauvegarde que les noms/métadonnées.
    """
    try:
        cleaned_data = []

        for phase in collected_data:
            clean_phase = {
                "phase_name": phase["phase_name"],
                "answers": {}
            }
            for k, v in phase["answers"].items():
                if isinstance(v, list) and v and hasattr(v[0], 'read'): 
                    file_names = ", ".join([f.name for f in v])
                    clean_phase["answers"][str(k)] = f"Fichiers (non stockés en DB): {file_names}"
                
                elif hasattr(v, 'read'): 
                     clean_phase["answers"][str(k)] = f"Fichier (non stocké en DB): {v.name}"
                else:
                    clean_phase["answers"][str(k)] = v
            
            cleaned_data.append(clean_phase)
        
        submission_id = st.session_state.get('submission_id', str(uuid.uuid4()))
        
        final_document = {
            "project_intitule": project_data.get('Intitulé', 'N/A'),
            "project_details": project_data,
            "submission_id": submission_id,
            "start_date": st.session_state.get('form_start_time', datetime.now()),
            "submission_date": datetime.now(),
            "status": "Completed",
            "collected_phases": cleaned_data
        }
        
        doc_id_base = str(project_data.get('Intitulé', 'form')).replace(" ", "_").replace("/", "_")[:20]
        doc_id = f"{doc_id_base}_{datetime.now().strftime('%Y%m%d_%H%M')}_{submission_id[:6]}"
        
        db.collection('FormAnswers').document(doc_id).set(final_document)
        return True, submission_id 
    except Exception as e:
        return False, str(e)

# --- FONCTIONS EXPORT CSV, ZIP ET WORD ---

def create_csv_export(collected_data, df_struct):
    rows = []
    submission_id = st.session_state.get('submission_id', 'N/A')
    project_name = st.session_state['project_data'].get('Intitulé', 'Projet Inconnu')
    start_time = st.session_state.get('form_start_time', 'N/A')
    end_time = datetime.now() 
    start_time_str = start_time.strftime('%Y-%m-%d %H:%M:%S') if isinstance(start_time, datetime) else 'N/A'
    end_time_str = end_time.strftime('%Y-%m-%d %H:%M:%S')

    for item in collected_data:
        phase_name = item['phase_name']
        for q_id, val in item['answers'].items():
            
            if int(q_id) == 100:
                q_text = "Commentaire Écart Photo"
            else:
                q_row = df_struct[df_struct['id'] == int(q_id)]
                q_text = q_row.iloc[0]['question'] if not q_row.empty else f"Question ID {q_id}"
            
            if isinstance(val, list) and val and hasattr(val[0], 'name'):
                final_val = f"[Pièces jointes] {len(val)} fichiers: " + ", ".join([f.name for f in val])
            elif hasattr(val, 'name'):
                final_val = f"[Pièce jointe] {val.name}"
            else:
                final_val = str(val)
            
            rows.append({
                "ID Formulaire": submission_id,
                "Date Début": start_time_str,
                "Date Fin": end_time_str,
                "Projet": project_name,
                "Phase": phase_name,
                "ID": q_id,
                "Question": q_text,
                "Réponse": final_val
            })
            
    df_export = pd.DataFrame(rows)
    return df_export.to_csv(index=False, sep=';', encoding='utf-8-sig')

def create_zip_export(collected_data):
    """
    Crée un ZIP contenant les photos présentes en mémoire.
    CORRECTION: Gestion correcte des fichiers UploadedFile de Streamlit
    """
    zip_buffer = io.BytesIO()
    
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        files_added = 0
        for phase in collected_data:
            phase_name_clean = str(phase['phase_name']).replace("/", "_").replace(" ", "_")
            
            for q_id, answer in phase['answers'].items():
                if isinstance(answer, list) and answer and hasattr(answer[0], 'read'):
                    for idx, file_obj in enumerate(answer):
                        try:
                            # CORRECTION: Reset du pointeur et lecture
                            file_obj.seek(0)
                            file_content = file_obj.read()
                            
                            # Vérification que le contenu n'est pas vide
                            if file_content:
                                original_name = file_obj.name.split('/')[-1].split('\\')[-1]
                                filename = f"{phase_name_clean}_Q{q_id}_{idx+1}_{original_name}"
                                zip_file.writestr(filename, file_content)
                                files_added += 1
                            
                            # Reset pour usage ultérieur
                            file_obj.seek(0)
                        except Exception as e:
                            st.warning(f"Erreur lors de l'ajout du fichier {file_obj.name}: {e}")
                            
        info_txt = f"Export généré le {datetime.now()}\nNombre de fichiers : {files_added}"
        zip_file.writestr("info.txt", info_txt)
    
    zip_buffer.seek(0)
    return zip_buffer


# --- NOUVELLES FONCTIONS UTILITAIRES POUR LA MISE EN FORME DES TABLES ---

def set_cell_borders(cell, **kwargs):
    """
    Définit les bordures d'une cellule spécifique.
    Les arguments sont des dictionnaires pour 'top', 'bottom', 'left', 'right', 
    chacun contenant 'color' (RGBColor), 'val' (bordure de style WD_BORDER), 'sz' (taille en points).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Propriétés de bordure par défaut
    default_border = {
        'val': 'single', # Style : simple ligne
        'sz': 4,         # Taille : 0.5 point (4 = 1/8 de point)
        'color': '000000' # Couleur : Noir
    }

    # Liste des bordures à modifier
    for key, value in kwargs.items():
        if key in ['top', 'bottom', 'left', 'right']:
            # Utilise un XMLElement pour la bordure
            tag = 'w:{}'.format(key)
            element = OxmlElement(tag)
            
            # Combinaison des valeurs par défaut et des valeurs fournies
            final_attrs = default_border.copy()
            if isinstance(value, dict):
                final_attrs.update(value)

            # Application des attributs
            element.set(qn('w:val'), final_attrs['val'])
            element.set(qn('w:sz'), str(final_attrs['sz']))
            element.set(qn('w:color'), final_attrs['color'])
            
            tcPr.append(element)

def set_cell_shading(cell, fill_color):
    """
    Définit la couleur de remplissage (ombrage) d'une cellule.
    fill_color doit être une chaîne HEX (ex: 'F0F0F0').
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:fill'), fill_color)
    tcPr.append(shading)


def define_custom_styles(doc):
    # ... (Styles de titre, sous-titre, texte restent inchangés) ...
    try:
        title_style = doc.styles.add_style('Report Title', WD_STYLE_TYPE.PARAGRAPH)
    except:
        title_style = doc.styles['Report Title']
    
    title_style.base_style = doc.styles['Heading 1']
    title_font = title_style.font
    title_font.name = 'Arial'
    title_font.size = Pt(20)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0x01, 0x38, 0x2D)# Bleu foncé
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(20)

    try:
        subtitle_style = doc.styles.add_style('Report Subtitle', WD_STYLE_TYPE.PARAGRAPH)
    except:
        subtitle_style = doc.styles['Report Subtitle']
        
    subtitle_style.base_style = doc.styles['Heading 2']
    subtitle_font = subtitle_style.font
    subtitle_font.name = 'Arial'
    subtitle_font.size = Pt(14)
    subtitle_font.bold = True
    subtitle_font.color.rgb = RGBColor(0x00, 0x56, 0x47) # Bleu moyen
    subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_style.paragraph_format.space_after = Pt(10)

    try:
        text_style = doc.styles.add_style('Report Text', WD_STYLE_TYPE.PARAGRAPH)
    except:
        text_style = doc.styles['Report Text']
        
    text_style.base_style = doc.styles['Normal']
    text_font = text_style.font
    text_font.name = 'Calibri'
    text_font.size = Pt(11)
    text_font.color.rgb = RGBColor(0x00, 0x00, 0x00) # Noir
    text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    text_style.paragraph_format.space_after = Pt(5)

    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    
    # 4. Style de Tableau (pour référence, la mise en forme physique se fera par cellule)
    try:
        doc.styles.add_style('Report Table Style', WD_STYLE_TYPE.TABLE)
    except:
        pass 
    
# --- Fonction principale modifiée ---

def create_word_report(collected_data, df_struct, project_data):
    """
    Crée un rapport Word avec toutes les questions et les photos
    """
    # ... (Début du document)
    doc = Document()
    define_custom_styles(doc) 
    
    doc.add_paragraph('Rapport d\'Audit Chantier', style='Report Title')
    doc.add_paragraph('Informations du Projet', style='Report Subtitle')
    
    # --- TABLEAU D'INFORMATIONS DU PROJET ---
    project_table = doc.add_table(rows=3, cols=2)
    
    # Centrage du tableau sur la page (Nouvel élément de dimension)
    project_table.allow_autofit = False # Important pour fixer les largeurs
    project_table.columns[0].width = Cm(5.0) # Largeur de la colonne d'en-tête
    project_table.columns[1].width = Cm(10.0) # Largeur de la colonne de valeur
    project_table.alignment = WD_ALIGN_TABLE.CENTER # Centrer le tableau entier

    # Bordure personnalisée (Bleu foncé, plus épaisse) pour les titres du tableau
    TITLE_BORDER = {
        'top': {'color': '01382D', 'sz': 8}, # 1 point
        'bottom': {'color': '01382D', 'sz': 8},
        'left': {'color': '01382D', 'sz': 8},
        'right': {'color': '01382D', 'sz': 8},
    }
    
    # Remplissage/Ombrage (Gris clair pour l'arrière-plan de l'en-tête)
    HEADER_FILL = 'F0F0F0' # Gris très clair
    
    # Remplir et formater le tableau
    project_data_items = [
        ('Intitulé', project_data.get('Intitulé', 'N/A')),
        ('Date de début', 'N/A'), # Sera remplacé plus bas
        ('Date de fin', 'N/A') # Sera remplacé plus bas
    ]
    
    try:
        # Les variables 'st.session_state' et 'datetime' doivent être définies ou importées
        start_time_str = st.session_state.get('form_start_time', datetime.now()).strftime('%d/%m/%Y %H:%M')
    except NameError:
        start_time_str = datetime.now().strftime('%d/%m/%Y %H:%M')
        
    project_data_items[1] = ('Date de début', start_time_str)
    project_data_items[2] = ('Date de fin', datetime.now().strftime('%d/%m/%Y %H:%M'))

    for i, (header, value) in enumerate(project_data_items):
        header_cell = project_table.rows[i].cells[0]
        value_cell = project_table.rows[i].cells[1]
        
        # 1. Mise à jour du texte
        header_cell.text = header
        value_cell.text = str(value)
        
        # 2. Application du style, alignement vertical et mise en gras
        for cell in [header_cell, value_cell]:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
            cell.paragraphs[0].style = 'Report Text'

        # Mise en forme spécifique de l'en-tête (colonne 0)
        header_cell.paragraphs[0].runs[0].bold = True
        
        # 3. Bordures et Remplissage
        if i == 0: # Ligne du haut, appliquer la bordure plus épaisse
             set_cell_borders(header_cell, **TITLE_BORDER)
             set_cell_borders(value_cell, **TITLE_BORDER)
        
        # Appliquer le remplissage gris clair à la colonne d'en-tête
        set_cell_shading(header_cell, HEADER_FILL)
        
    # Appliquer une bordure simple aux autres lignes si vous ne voulez pas du style 'Light Grid'
    # Sinon, vous pouvez utiliser project_table.style = 'Light Grid' et supprimer les appels à set_cell_borders.
    
    # ... (Reste du code d'information du projet inchangé)
    doc.add_paragraph()
    doc.add_paragraph('Détails du Projet', style='Report Subtitle')
    
    # ... (Boucle des détails du projet inchangée) ...
    # ...
    
    doc.add_page_break()
    
    # Parcourir toutes les phases
    for phase_idx, phase in enumerate(collected_data):
        # ... (En-tête de phase inchangé) ...
        doc.add_paragraph(f"Phase: {phase['phase_name']}", style='Report Subtitle')
        
        # Parcourir toutes les questions de cette phase
        for q_id, answer in phase['answers'].items():
            
            # ... (Récupération de q_text et vérification is_photo_answer inchangés) ...
            
            if not is_photo_answer:
                # --- Affichage des autres réponses sous forme de tableau ---
                
                table = doc.add_table(rows=1, cols=2)
                
                # 1. Dimensions et Style du tableau de question
                table.allow_autofit = False # Fixer la taille manuellement
                table.alignment = WD_ALIGN_TABLE.LEFT # Aligner à gauche
                table.style = 'Table Grid' # Un style simple pour commencer

                # Largeurs des colonnes (Question 80%, Réponse 20%)
                q_cell = table.cell(0, 0)
                a_cell = table.cell(0, 1)
                
                q_cell.width = Inches(5.5) 
                a_cell.width = Inches(1.5) 
                
                # 2. Bordures personnalisées (Vert foncé)
                QUESTION_BORDER = {
                    'top': {'color': '005647', 'sz': 6}, # 0.75 point
                    'bottom': {'color': '005647', 'sz': 6}, 
                    'left': {'color': '005647', 'sz': 6},
                    'right': {'color': '005647', 'sz': 6},
                }
                
                # Bordure intérieure (plus fine ou différente)
                MIDDLE_BORDER = {
                    'right': {'color': '005647', 'sz': 6}
                }

                # Appliquer la bordure à la cellule Question
                set_cell_borders(q_cell, **QUESTION_BORDER)
                # Appliquer la bordure à la cellule Réponse
                set_cell_borders(a_cell, **QUESTION_BORDER)
                
                # 3. Remplissage (Gris clair pour la question)
                set_cell_shading(q_cell, HEADER_FILL)
                
                # 4. Contenu et styles
                q_cell.text = f'Q{q_id}: {q_text}'
                a_cell.text = str(answer)

                for cell in table.rows[0].cells:
                    cell.paragraphs[0].style = 'Report Text'
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                    
                q_cell.paragraphs[0].runs[0].bold = True
                
                # Centrer la réponse horizontalement
                a_cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                doc.add_paragraph() # Espace entre les tableaux/questions
            
            # ... (Le code d'affichage des photos reste inchangé) ...
        
        # Saut de page entre les phases (sauf pour la dernière)
        if phase_idx < len(collected_data) - 1:
            doc.add_page_break()
    
    # Sauvegarder dans un buffer
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)
    
    return word_buffer

# --- GESTION DE L'ÉTAT (inchangée) ---
def init_session_state():
    defaults = {
        'step': 'PROJECT_LOAD',
        'project_data': None,
        'collected_data': [],
        'current_phase_temp': {},
        'current_phase_name': None,
        'iteration_id': str(uuid.uuid4()), 
        'identification_completed': False,
        'data_saved': False,
        'id_rendering_ident': None,
        'form_start_time': None,
        'submission_id': None,
        'show_comment_on_error': False
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value

init_session_state()

# --- LOGIQUE MÉTIER (inchangée) ---

def check_condition(row, current_answers, collected_data):
    try:
        if int(row.get('Condition on', 0)) != 1: return True
    except (ValueError, TypeError): return True

    all_past_answers = {}
    for phase_data in collected_data: all_past_answers.update(phase_data['answers'])
    combined_answers = {**all_past_answers, **current_answers}
    
    condition_str = str(row.get('Condition value', '')).strip()
    if not condition_str or "=" not in condition_str: return True

    try:
        target_id_str, expected_value_raw = condition_str.split('=', 1)
        target_id = int(target_id_str.strip())
        expected_value = expected_value_raw.strip().strip('"').strip("'")
        user_answer = combined_answers.get(target_id)
        if user_answer is not None:
            return str(user_answer).lower() == str(expected_value).lower()
        else:
            return False
    except Exception: return True

# -----------------------------------------------------------
# --- FONCTION VALIDATION (Identique) ---
# -----------------------------------------------------------
COMMENT_ID = 100
COMMENT_QUESTION = "Veuillez préciser pourquoi le nombre de photo partagé ne correspond pas au minimum attendu"

def validate_section(df_questions, section_name, answers, collected_data):
    missing = []
    section_rows = df_questions[df_questions['section'] == section_name]
    
    comment_val = answers.get(COMMENT_ID)
    has_justification = comment_val is not None and str(comment_val).strip() != ""
    project_data = st.session_state.get('project_data', {})
    
    expected_total_base, detail_str = get_expected_photo_count(section_name.strip(), project_data)
    expected_total = expected_total_base
    
    photo_question_count = sum(
        1 for _, row in section_rows.iterrows()
        if str(row.get('type', '')).strip().lower() == 'photo' and check_condition(row, answers, collected_data)
    )
    
    if expected_total is not None and expected_total > 0:
        expected_total = expected_total_base * photo_question_count
        detail_str = (
            f"{detail_str} | Questions photo visibles: {photo_question_count} "
            f"-> Total ajusté: {expected_total}"
        )

    current_photo_count = 0
    photo_questions_found = False
    
    for _, row in section_rows.iterrows():
        q_type = str(row['type']).strip().lower()
        if q_type == 'photo' and check_condition(row, answers, collected_data):
            photo_questions_found = True
            q_id = int(row['id'])
            val = answers.get(q_id)
            if isinstance(val, list):
                current_photo_count += len(val)

    is_count_sufficient = (
        expected_total is None or expected_total <= 0 or 
        (expected_total > 0 and current_photo_count >= expected_total)
    )
    
    for _, row in section_rows.iterrows():
        if int(row['id']) == COMMENT_ID: continue
        if not check_condition(row, answers, collected_data): continue
        
        is_mandatory = str(row['obligatoire']).strip().lower() == 'oui'
        q_id = int(row['id'])
        q_type = str(row['type']).strip().lower()
        val = answers.get(q_id)
        
        if is_mandatory:
            if q_type == 'photo':
                if is_count_sufficient or has_justification:
                    continue
                else:
                    pass

            if isinstance(val, list):
                if not val: missing.append(f"Question {q_id} : {row['question']} (fichier(s) manquant(s))")
            elif val is None or val == "" or (isinstance(val, (int, float)) and val == 0):
                missing.append(f"Question {q_id} : {row['question']}")

    is_photo_count_incorrect = False
    if expected_total is not None and expected_total > 0:
        if photo_questions_found and current_photo_count != expected_total:
            is_photo_count_incorrect = True
            error_message = (
                f"⚠️ **Écart de Photos pour '{str(section_name)}'**.\n\n"
                f"Attendu : **{str(expected_total)}** (calculé : {str(detail_str)}).\n"
                f"Reçu : **{str(current_photo_count)}**.\n\n"
                f"Le champ de commentaire doit être rempli."
            )
            if not has_justification:
                missing.append(
                    f"**Commentaire (ID {COMMENT_ID}) :** {COMMENT_QUESTION} "
                    f"(requis en raison de l'écart de photo : Attendu {expected_total}, Reçu {current_photo_count}).\n\n"
                    f"{error_message}"
                )

    if not is_photo_count_incorrect and COMMENT_ID in answers:
        del answers[COMMENT_ID]

    return len(missing) == 0, missing

validate_phase = validate_section
validate_identification = validate_section

# --- COMPOSANTS UI (inchangés) ---

def render_question(row, answers, phase_name, key_suffix, loop_index):
    q_id = int(row.get('id', 0))
    is_dynamic_comment = q_id == COMMENT_ID
    if is_dynamic_comment:
        q_text = COMMENT_QUESTION
        q_type = 'text' 
        q_desc = "Ce champ est obligatoire si le nombre de photos n'est pas conforme."
        q_mandatory = True 
        q_options = []
    else:
        q_text = row['question']
        q_type = str(row['type']).strip().lower()
        q_desc = row['Description']
        q_mandatory = str(row['obligatoire']).lower() == 'oui'
        q_options = str(row['options']).split(',') if row['options'] else []
        
    q_text = str(q_text).strip()
    q_desc = str(q_desc).strip()
    label_html = f"<strong>{q_id}. {q_text}</strong>" + (' <span class="mandatory">*</span>' if q_mandatory else "")
    widget_key = f"q_{q_id}_{phase_name}_{key_suffix}_{loop_index}"
    current_val = answers.get(q_id)
    val = current_val

    st.markdown(f'<div class="question-card"><div>{label_html}</div>', unsafe_allow_html=True)
    if q_desc: st.markdown(f'<div class="description">⚠️ {q_desc}</div>', unsafe_allow_html=True)

    if q_type == 'text':
        if is_dynamic_comment:
             val = st.text_area("Justification de l'écart", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed")
        else:
             val = st.text_input("Réponse", value=current_val if current_val else "", key=widget_key, label_visibility="collapsed")

    elif q_type == 'select':
        clean_opts = [opt.strip() for opt in q_options]
        if "" not in clean_opts: clean_opts.insert(0, "")
        idx = clean_opts.index(current_val) if current_val in clean_opts else 0
        val = st.selectbox("Sélection", clean_opts, index=idx, key=widget_key, label_visibility="collapsed")
    
    elif q_type == 'number':
        if q_id == 9:
            default_val = int(float(current_val)) if current_val is not None and str(current_val).replace('.', '', 1).isdigit() else 0
            val = st.number_input("Nombre (entier)", value=default_val, step=1, format="%d", key=widget_key, label_visibility="collapsed")
        else:
            default_val = float(current_val) if current_val and str(current_val).replace('.', '', 1).isdigit() else 0.0
            val = st.number_input("Nombre", value=default_val, key=widget_key, label_visibility="collapsed")
    
    elif q_type == 'photo':
        expected, details = get_expected_photo_count(phase_name.strip(), st.session_state.get('project_data'))
        if expected is not None and expected > 0:
            st.info(f"📸 **Photos :** Il est attendu **{expected}** photos pour cette section (Base calculée : {details}).")
            st.divider()

        val = st.file_uploader("Images", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key=widget_key, label_visibility="collapsed")
        
        if val:
            file_names = ", ".join([f.name for f in val])
            st.success(f"Nombre d'images chargées : {len(val)} ({file_names})")
        elif current_val and isinstance(current_val, list) and current_val:
            names = ", ".join([getattr(f, 'name', 'Fichier') for f in current_val])
            st.info(f"Fichiers conservés : {len(current_val)} ({names})")
    
    st.markdown('</div>', unsafe_allow_html=True)
    
    if val is not None and (not is_dynamic_comment or val.strip() != ""): answers[q_id] = val 
    elif current_val is not None and not is_dynamic_comment: answers[q_id] = current_val
    elif is_dynamic_comment and (val is None or val.strip() == ""):
        if q_id in answers: del answers[q_id]

# --- FLUX PRINCIPAL ---

st.markdown('<div class="main-header"><h1>📝Formulaire Chantier </h1></div>', unsafe_allow_html=True)

if st.session_state['step'] == 'PROJECT_LOAD':
    st.info("Tentative de chargement de la structure des formulaires...")
    with st.spinner("Chargement en cours..."):
        df_struct = load_form_structure_from_firestore()
        df_site = load_site_data_from_firestore()
        
        if df_struct is not None and df_site is not None:
            st.session_state['df_struct'] = df_struct
            st.session_state['df_site'] = df_site
            st.session_state['step'] = 'PROJECT'
            st.rerun()
        else:
            st.error("Impossible de charger les données. Vérifiez votre connexion et les secrets Firebase.")
            if st.button("Réessayer le chargement"):
                load_form_structure_from_firestore.clear() 
                load_site_data_from_firestore.clear() 
                st.session_state['step'] = 'PROJECT_LOAD'
                st.rerun()

elif st.session_state['step'] == 'PROJECT':
    df_site = st.session_state['df_site']
    st.markdown("### 🏗️ Sélection du Chantier")
    
    if 'Intitulé' not in df_site.columns:
        st.error("Colonne 'Intitulé' manquante dans les données 'Sites'.")
    else:
        search_term = st.text_input("Rechercher un projet (Veuillez renseigner au minimum 3 caractères pour le nom de la ville)", key="project_search_input").strip()
        filtered_projects = []
        selected_proj = None
        
        if len(search_term) >= 3:
            mask = df_site['Intitulé'].str.contains(search_term, case=False, na=False)
            filtered_projects_df = df_site[mask]
            filtered_projects = [""] + filtered_projects_df['Intitulé'].dropna().unique().tolist()
            if filtered_projects:
                selected_proj = st.selectbox("Résultats de la recherche", filtered_projects)
            else:
                st.warning(f"Aucun projet trouvé pour **'{search_term}'**.")
        elif len(search_term) > 0 and len(search_term) < 3:
            st.info("Veuillez entrer au moins **3 caractères** pour lancer la recherche.")
        
        if selected_proj:
            row = df_site[df_site['Intitulé'] == selected_proj].iloc[0]
            st.info(f"Projet sélectionné : **{selected_proj}**")
            if st.button("✅ Démarrer l'identification"):
                st.session_state['project_data'] = row.to_dict()
                st.session_state['form_start_time'] = datetime.now() 
                st.session_state['submission_id'] = str(uuid.uuid4())
                st.session_state['step'] = 'IDENTIFICATION'
                st.session_state['current_phase_temp'] = {}
                st.session_state['iteration_id'] = str(uuid.uuid4())
                st.session_state['show_comment_on_error'] = False
                st.rerun()

elif st.session_state['step'] == 'IDENTIFICATION':
    df = st.session_state['df_struct']
    ID_SECTION_NAME = df['section'].iloc[0]
    st.markdown(f"### 👤 Étape unique : {ID_SECTION_NAME}")
    identification_questions = df[df['section'] == ID_SECTION_NAME]
    if st.session_state['id_rendering_ident'] is None: st.session_state['id_rendering_ident'] = str(uuid.uuid4())
    rendering_id = st.session_state['id_rendering_ident']
    
    for idx, (index, row) in enumerate(identification_questions.iterrows()):
        if check_condition(row, st.session_state['current_phase_temp'], st.session_state['collected_data']):
            render_question(row, st.session_state['current_phase_temp'], ID_SECTION_NAME, rendering_id, idx)
            
    st.markdown("---")
    if st.button("✅ Valider l'identification"):
        is_valid, errors = validate_identification(df, ID_SECTION_NAME, st.session_state['current_phase_temp'], st.session_state['collected_data'])
        if is_valid:
            id_entry = {"phase_name": ID_SECTION_NAME, "answers": st.session_state['current_phase_temp'].copy()}
            st.session_state['collected_data'].append(id_entry)
            st.session_state['identification_completed'] = True
            st.session_state['step'] = 'LOOP_DECISION'
            st.session_state['current_phase_temp'] = {}
            st.session_state['show_comment_on_error'] = False
            st.success("Identification validée.")
            st.rerun()
        else:
            st.markdown('<div class="error-box"><b>⚠️ Erreur de validation :</b><br>' + '<br>'.join([f"- {e}" for e in errors]) + '</div>', unsafe_allow_html=True)

elif st.session_state['step'] in ['LOOP_DECISION', 'FILL_PHASE']:
    project_intitule = st.session_state['project_data'].get('Intitulé', 'Projet Inconnu')
    with st.expander(f"📍 Projet : {project_intitule}", expanded=False):
        project_details = st.session_state['project_data']
        st.markdown(":orange-badge[**Détails du Projet sélectionné :**]")
        
        with st.container(border=True):
            st.markdown("**Informations générales**")
            cols1 = st.columns([1, 1, 1]) 
            fields_l1 = DISPLAY_GROUPS[0]
            for i, field_key in enumerate(fields_l1):
                renamed_key = PROJECT_RENAME_MAP.get(field_key, field_key)
                value = project_details.get(field_key, 'N/A')
                with cols1[i]: st.markdown(f"**{renamed_key}** : {value}")
                    
        with st.container(border=True):
            st.markdown("**Points de charge Standard**")
            cols2 = st.columns([1, 1, 1])
            fields_l2 = DISPLAY_GROUPS[1]
            for i, field_key in enumerate(fields_l2):
                renamed_key = PROJECT_RENAME_MAP.get(field_key, field_key)
                value = project_details.get(field_key, 'N/A')
                with cols2[i]: st.markdown(f"**{renamed_key}** : {value}")

        with st.container(border=True):
            st.markdown("**Points de charge Pré-équipés**")
            cols3 = st.columns([1, 1, 1])
            fields_l3 = DISPLAY_GROUPS[2]
            for i, field_key in enumerate(fields_l3):
                renamed_key = PROJECT_RENAME_MAP.get(field_key, field_key)
                value = project_details.get(field_key, 'N/A')
                with cols3[i]: st.markdown(f"**{renamed_key}** : {value}")
        
        st.write(":orange-badge[**Phases et Identification déjà complétées :**]")
        for idx, item in enumerate(st.session_state['collected_data']):
            st.write(f"• **{item['phase_name']}** : {len(item['answers'])} réponses")

    if st.session_state['step'] == 'LOOP_DECISION':
        st.markdown("### 🔄 Gestion des Phases")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("➕ Ajouter une phase"):
                st.session_state['step'] = 'FILL_PHASE'
                st.session_state['current_phase_temp'] = {}
                st.session_state['current_phase_name'] = None
                st.session_state['iteration_id'] = str(uuid.uuid4())
                st.session_state['show_comment_on_error'] = False
                st.rerun()
        with col2:
            if st.button("🏁 Terminer l'audit"):
                st.session_state['step'] = 'FINISHED'
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)

    elif st.session_state['step'] == 'FILL_PHASE':
        df = st.session_state['df_struct']
        ID_SECTION_NAME = df['section'].iloc[0]
        ID_SECTION_CLEAN = str(ID_SECTION_NAME).strip().lower()
        SECTIONS_TO_EXCLUDE_CLEAN = {ID_SECTION_CLEAN, "phase"}
        all_sections_raw = df['section'].unique().tolist()
        available_phases = []
        for sec in all_sections_raw:
            if pd.isna(sec) or not sec or str(sec).strip().lower() in SECTIONS_TO_EXCLUDE_CLEAN: continue
            available_phases.append(sec)
        
        if not st.session_state['current_phase_name']:
              st.markdown("### 📑 Sélection de la phase")
              phase_choice = st.selectbox("Quelle phase ?", [""] + available_phases)
              if phase_choice:
                  st.session_state['current_phase_name'] = phase_choice
                  st.session_state['show_comment_on_error'] = False 
                  st.rerun()
              if st.button("⬅️ Retour"):
                  st.session_state['step'] = 'LOOP_DECISION'
                  st.session_state['current_phase_temp'] = {}
                  st.session_state['show_comment_on_error'] = False
                  st.rerun()
        else:
            current_phase = st.session_state['current_phase_name']
            st.markdown(f"### 📝 {current_phase}")
            if st.button("🔄 Changer de phase"):
                st.session_state['current_phase_name'] = None
                st.session_state['current_phase_temp'] = {}
                st.session_state['iteration_id'] = str(uuid.uuid4())
                st.session_state['show_comment_on_error'] = False 
                st.rerun()
            st.divider()
            
            section_questions = df[df['section'] == current_phase]
            visible_count = 0
            for idx, (index, row) in enumerate(section_questions.iterrows()):
                if int(row.get('id', 0)) == COMMENT_ID: continue
                if check_condition(row, st.session_state['current_phase_temp'], st.session_state['collected_data']):
                    render_question(row, st.session_state['current_phase_temp'], current_phase, st.session_state['iteration_id'], idx)
                    visible_count += 1
            
            if visible_count == 0 and not st.session_state.get('show_comment_on_error', False):
                st.warning("Aucune question visible.")

            if st.session_state.get('show_comment_on_error', False):
                st.markdown("---")
                st.markdown("### ✍️ Justification de l'Écart")
                comment_row = pd.Series({'id': COMMENT_ID})
                render_question(comment_row, st.session_state['current_phase_temp'], current_phase, st.session_state['iteration_id'], 999) 
            
            st.markdown("---")
            c1, c2 = st.columns([1, 2])
            with c1:
                if st.button("❌ Annuler"):
                    st.session_state['step'] = 'LOOP_DECISION'
                    st.session_state['show_comment_on_error'] = False
                    st.rerun()
            with c2:
                if st.button("💾 Valider la phase"):
                    st.session_state['show_comment_on_error'] = False 
                    is_valid, errors = validate_phase(df, current_phase, st.session_state['current_phase_temp'], st.session_state['collected_data'])
                    if is_valid:
                        new_entry = {"phase_name": current_phase, "answers": st.session_state['current_phase_temp'].copy()}
                        st.session_state['collected_data'].append(new_entry)
                        st.success("Enregistré !")
                        st.session_state['step'] = 'LOOP_DECISION'
                        st.rerun()
                    else:
                        is_photo_error = any(f"Commentaire (ID {COMMENT_ID})" in e for e in errors)
                        if is_photo_error: st.session_state['show_comment_on_error'] = True
                        html_errors = '<br>'.join([f"- {e}" for e in errors])
                        st.markdown(f'<div class="error-box"><b>⚠️ Erreurs :</b><br>{html_errors}</div>', unsafe_allow_html=True)
                        st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

elif st.session_state['step'] == 'FINISHED':
    st.markdown("## 🎉 Formulaire Terminé")
    project_name = st.session_state['project_data'].get('Intitulé', 'Projet Inconnu')
    st.write(f"Projet : **{project_name}**")
    
    # 1. SAUVEGARDE FIREBASE
    if not st.session_state['data_saved']:
        with st.spinner("Sauvegarde des réponses dans Firestore..."):
            success, submission_id_returned = save_form_data(
                st.session_state['collected_data'], 
                st.session_state['project_data']
            )

            if success:
                st.success(f"Données textuelles sauvegardées sur Firestore ! (ID: {submission_id_returned})")
                st.session_state['data_saved'] = True
            else:
                st.error(f"Erreur lors de la sauvegarde : {submission_id_returned}")
                if st.button("Réessayer la sauvegarde"):
                    st.rerun()
    else:
        st.info("Les données ont déjà été sauvegardées sur Firestore.")

    st.markdown("---")
    
    if st.session_state['data_saved']:
        # Préparation des exports
        csv_data = create_csv_export(st.session_state['collected_data'], st.session_state['df_struct'])
        zip_buffer = create_zip_export(st.session_state['collected_data'])
        date_str = datetime.now().strftime('%Y%m%d_%H%M')
        
        # --- 2. TÉLÉCHARGEMENT DIRECT ---
        st.markdown("### 📥 Télécharger les fichiers")
        
        col_csv, col_zip, col_word = st.columns(3)
        
        file_name_csv = f"Export_{project_name}_{date_str}.csv"
        with col_csv:
            st.download_button(
                label="📄 CSV", 
                data=csv_data, 
                file_name=file_name_csv, 
                mime='text/csv',
                use_container_width=True
            )

        if zip_buffer:
            file_name_zip = f"Photos_{project_name}_{date_str}.zip"
            with col_zip:
                st.download_button(
                    label="📸 ZIP Photos", 
                    data=zip_buffer.getvalue(), 
                    file_name=file_name_zip, 
                    mime='application/zip',
                    use_container_width=True
                )
        
        # Génération du rapport Word
        with st.spinner("Génération du rapport Word..."):
            try:
                word_buffer = create_word_report(
                    st.session_state['collected_data'],
                    st.session_state['df_struct'],
                    st.session_state['project_data']
                )
                
                file_name_word = f"Rapport_{project_name}_{date_str}.docx"
                with col_word:
                    st.download_button(
                        label="📋 Rapport Word", 
                        data=word_buffer.getvalue(), 
                        file_name=file_name_word, 
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        use_container_width=True
                    )
                st.success("✅ Rapport Word généré avec succès !")
            except Exception as e:
                st.error(f"Erreur lors de la génération du rapport Word : {e}")
    
        # --- 3. OUVERTURE DE L'APPLICATION NATIVE (MAILTO) ---
        st.markdown("---")
        st.markdown("### 📧 Partager par Email")
        st.info("💡 Téléchargez d'abord les fichiers ci-dessus, puis cliquez sur le bouton ci-dessous pour ouvrir votre application email.")
        
        subject = f"Rapport Audit : {project_name}"
        body = (
            f"Bonjour,\n\n"
            f"Veuillez trouver ci-joint le rapport d'audit pour le projet {project_name}.\n"
            f"Fichiers à joindre :\n"
            f"- {file_name_csv}\n"
            f"- {file_name_zip}\n"
            f"- {file_name_word}\n\n"
            f"Cordialement."
        )
        
        mailto_link = (
            f"mailto:?" 
            f"subject={urllib.parse.quote(subject)}" 
            f"&body={urllib.parse.quote(body)}"
        )
        
        st.markdown(
            f'<a href="{mailto_link}" target="_blank" style="text-decoration: none;">'
            f'<button style="background-color: #E9630C; color: white; border: none; padding: 10px 20px; border-radius: 8px; width: 100%; font-size: 16px; cursor: pointer;">'
            f'📧 Ouvrir l\'application Email'
            f'</button>'
            f'</a>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    if st.button("🔄 Recommencer l'audit"):
        st.session_state.clear()
        st.rerun()
    
